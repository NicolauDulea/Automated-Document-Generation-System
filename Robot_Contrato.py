import PySimpleGUI as sg
import pandas as pd
from docx import Document
from datetime import datetime
import os  
from docx.shared import Pt
sg.theme("reddit")

layout = [
    [sg.Text("Excel"), 
     sg.FileBrowse("Escolher Ficheiro Excel", target="input_excel", button_color=('white', 'green')), 
     sg.Input(key="input_excel")],
    [sg.Text("Word"), 
     sg.FileBrowse("Escolher Ficheiro Word", target="input_word"), 
     sg.Input(key="input_word")],
    [sg.Button("Gerar Contratos Completos")]
]



window = sg.Window("Gerador de Contrato", layout)

def preencher_contrato(excel_path, word_path, folder_name, today):
    
    try:
        excel = pd.read_excel(excel_path)

        # Iterate over each row in the Excel file
        for _, row in excel.iterrows():
            word_doc = Document(word_path)
                
            nome = row.iloc[3]
            estado_civil = row.iloc[4]
            morada = f"{row.iloc[5]}, {row.iloc[6]}" 
            tipo_id = str(row.iloc[7])
            nr_id = str(row.iloc[8])
            
            
            validade_raw = str(row.iloc[9])

            if " / " in validade_raw:
                # Split and remove any trailing empty parts
                validade_list = [date.strip() for date in validade_raw.split(" / ") if date.strip()]
                validade_dates = [
                    pd.to_datetime(date, errors='coerce') for date in validade_list
                ]
                validade_dates = [date for date in validade_dates if not pd.isna(date)]
                pre_validade = max(validade_dates) if validade_dates else None
            else:
                # Single date processing
                pre_validade = pd.to_datetime(validade_raw.strip(), errors='coerce')
            
            validade = pre_validade.strftime('%d/%m/%Y')
            nif = row.iloc[11]

            # Prepare the text to insert
            replacement_text = (
                f"{nome}, {estado_civil}, residente na {morada}, "
                f"contribuinte fiscal n.º {nif}, portador do {tipo_id} n.º {nr_id}, "
                f"válido até {validade}, e do NISS    , de ora em diante designado apenas por “o Trabalhador”."
            )

            # Replace the placeholder text in the document
            for paragraph in word_doc.paragraphs:
                if "portador do Cartão de Cidadão" in paragraph.text:
                    paragraph.clear()
                    run = paragraph.add_run(replacement_text)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)


                # Replace the date in the format "Lisboa, no dia ... de … de 202…." or "Lisboa, ... de … de 202…."
                if "Lisboa, no dia" in paragraph.text or "Lisboa," in paragraph.text:
                    # Replace with today's date
                    paragraph.text = paragraph.text.replace("Lisboa, no dia ... de … de 202…", f"Lisboa, no dia {today}")
                    paragraph.text = paragraph.text.replace("Lisboa, ... de … de 202…", f"Lisboa, {today}")
                      # Set the font to Calibri for the new date text
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)  

            # Save the modified Word document in the new folder with a unique name for each row
            output_path = os.path.join(folder_name, f"{nome}.docx")  # Save each file in the new folder
            word_doc.save(output_path)
            print(f"Documento gerado: {output_path}")
            
        sg.popup("Documentos gerados com sucesso!", "Arquivos foram salvos em uma nova pasta.")
    
    
    except Exception as e:
        sg.popup("Erro ao ler o arquivo Excel:", str(e))
        return


# Loop da interface para pegar os arquivos
while True:
    event, values = window.read() # Listen for events
    if event == sg.WIN_CLOSED:     # Check if the window was closed
        break                       # Exit the loop
    elif event == "Gerar Contratos Completos":          # Check for other events
        today = datetime.today().strftime("%d de %B de %Y")
        desktop_path = os.path.expanduser("~/Desktop")
        folder_name = os.path.join(desktop_path, f"Contratos - {today}")

        os.makedirs(folder_name, exist_ok=True)

        
        excel_path = values["input_excel"]
        word_path = values["input_word"]
        
        # Chama a função para preencher o contrato
        preencher_contrato(excel_path, word_path, folder_name, today)

window.close()

